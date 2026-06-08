#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Парсер сводки ЕДДС МКУ «СолнСпас» (docx) -> структурированный JSON.
Рассчитан на типовой шаблон сводки. Устойчив к пустым/частично заполненным таблицам.
"""
import sys, os, json, re
import docx


def clean(s):
    if s is None:
        return ""
    return re.sub(r'[ \t]+', ' ', s.replace('\xa0', ' ')).strip()


def cell_texts_dedup(row):
    """Возвращает уникальные значения ячеек строки по порядку (в шаблоне ячейки
    часто объединены и дублируются из-за merge)."""
    out = []
    for c in row.cells:
        t = clean(c.text)
        if not out or out[-1] != t:
            out.append(t)
    return out


def first_meaningful(row):
    """Первое непустое уникальное значение в строке (для merged-описаний)."""
    seen = set()
    vals = []
    for c in row.cells:
        t = clean(c.text)
        if t and t not in seen:
            seen.add(t)
            vals.append(t)
    return vals


def split_events(text):
    """Разбивает длинное описание происшествий на отдельные эпизоды по датам."""
    if not text:
        return []
    # эпизоды начинаются с даты вида 31.05.2026г.
    parts = re.split(r'(?=\d{2}\.\d{2}\.\d{4}\s*г?\.?\s*(?:в\s*)?\d{1,2}[:.]\d{2}|(?<!\d)\d{2}\.\d{2}\.\d{4}г\.)', text)
    parts = [clean(p) for p in parts if clean(p)]
    return parts if parts else [clean(text)]


def parse(path):
    d = docx.Document(path)

    data = {
        "meta": {},
        "header": {},
        "incidents": [],          # таблица происшествий/преступлений
        "fire": {},               # пожарная обстановка
        "appeals": {},            # обращения граждан
        "tech_violations": [],    # технологические нарушения (ЦО/ГВС/ХВС/Газ/...)
        "hotline_categories": [], # горячая линия Главы по категориям
        "edds_stats": {},         # статистика вызовов ЕДДС
        "spas_stats": {},         # статистика выездов СОЛН СПАС
        "water_levels": [],       # уровни воды
        "weather": {},            # погода
        "warnings": [],           # предупреждения о НГЯ
        "fire_danger_class": "",
        "signed_by": "",
    }

    # --- Заголовок из параграфов ---
    paras = [clean(p.text) for p in d.paragraphs if clean(p.text)]
    full_text = "\n".join(paras)

    # Дата периода
    m = re.search(r'с\s+(\d{1,2}\s+\w+)\s+на\s+(\d{1,2}\s+\w+\s+\d{4})', full_text)
    if m:
        data["header"]["period"] = f"с {m.group(1)} на {m.group(2)}"
    m = re.search(r'по состоянию на\s*(\d{1,2}\s*ч\s*\d{1,2}\s*мин)', full_text)
    if m:
        data["header"]["as_of"] = clean(m.group(1))
    # Дата отчёта.
    # ПРИОРИТЕТ 1: из имени файла "Svodka-ДД.ММ.ГГГГ" — это официальная дата смены ЕДДС.
    # Ночная смена работает до 8 утра, поэтому в тексте справки («на 8 июня»)
    # может стоять дата следующего дня — её НЕ используем как дату отчёта.
    # ПРИОРИТЕТ 2 (fallback): дата из текста «на <число> <месяц> <год>».
    months = {"января":1,"февраля":2,"марта":3,"апреля":4,"мая":5,"июня":6,
              "июля":7,"августа":8,"сентября":9,"октября":10,"ноября":11,"декабря":12}
    fname = os.path.basename(path)
    mf = re.search(r'(\d{1,2})[.\-_](\d{1,2})[.\-_](\d{4})', fname)
    if mf:
        data["meta"]["date"] = f"{mf.group(3)}-{int(mf.group(2)):02d}-{int(mf.group(1)):02d}"
    else:
        m = re.search(r'на\s+(\d{1,2})\s+(\w+)\s+(\d{4})', full_text)
        if m and m.group(2).lower() in months:
            data["meta"]["date"] = f"{m.group(3)}-{months[m.group(2).lower()]:02d}-{int(m.group(1)):02d}"

    data["header"]["org"] = "МКУ «Аварийно-спасательная служба городского округа Солнечногорск Московской области»"
    if "чрезвычайных ситуаций и крупных аварий не зафиксировано" in full_text:
        data["header"]["cs_status"] = "За прошедшие сутки на территории городского округа чрезвычайных ситуаций и крупных аварий не зафиксировано."

    # Предупреждения о НГЯ
    for m in re.finditer(r'Предупреждение о НГЯ\s*№\s*(\d+)\s*(.*?)(?=Предупреждение о НГЯ|Старший дежурный|$)', full_text, re.S):
        num = m.group(1)
        body = clean(re.sub(r'\s+', ' ', m.group(2)))
        data["warnings"].append({"num": num, "text": body})

    m = re.search(r'(\w+)\s+класса пожарной опасности', full_text)
    if m:
        mm = re.search(r'прогнозируется преобладание\s+(.*?класса пожарной опасности)', full_text)
        if mm:
            data["fire_danger_class"] = clean(mm.group(0))

    m = re.search(r'(Опасные метеорологические явления:.*)', full_text)
    if m:
        data["weather"]["dangerous"] = clean(m.group(1).split('\n')[0])
    m = re.search(r'(Неблагоприятные метеорологические явления:.*)', full_text)
    if m:
        data["weather"]["unfavorable"] = clean(m.group(1).split('\n')[0])

    m = re.search(r'(С\.А\.\s*\w+|[А-Я]\.[А-Я]\.\s*\w+)\.?\s*$', full_text)
    if m:
        data["signed_by"] = clean(m.group(1))

    tables = d.tables

    def find_table(keyword, col=0):
        for t in tables:
            for r in t.rows:
                if r.cells and keyword.lower() in clean(r.cells[0].text).lower():
                    return t
        return None

    # --- Таблица происшествий ---
    for t in tables:
        head = clean(t.rows[0].cells[0].text).lower() if t.rows else ""
        if "вид происшествия" in head:
            for r in t.rows[1:]:
                vals = cell_texts_dedup(r)
                if len(vals) < 2:
                    continue
                name = vals[0].rstrip(':')
                count = vals[1] if len(vals) > 1 else ""
                desc = vals[2] if len(vals) > 2 else ""
                events = split_events(desc) if desc else []
                data["incidents"].append({
                    "type": name, "count": count, "events": events
                })
            break

    # --- Пожарная обстановка ---
    for t in tables:
        if t.rows and "кто" in clean(t.rows[0].cells[0].text).lower() and "выезжал" in clean(t.rows[0].cells[0].text).lower():
            # последняя строка — значения
            vals = cell_texts_dedup(t.rows[-1])
            data["fire"]["raw"] = vals
            all_dash = all(v in ('-', '', '–') for v in vals)
            data["fire"]["empty"] = all_dash
            break

    # --- Обращения граждан (большая таблица с горячей линией) ---
    for ti, t in enumerate(tables):
        if t.rows and "горячая линия" in clean(t.rows[0].cells[0].text).lower():
            # строка 0 — заголовки категорий, строка 1 — числа
            cats = cell_texts_dedup(t.rows[0])
            nums = cell_texts_dedup(t.rows[1]) if len(t.rows) > 1 else []
            summary = {}
            for k, v in zip(cats, nums):
                summary[k] = v
            data["appeals"]["summary"] = summary

            # Технологические нарушения: блоки по типу ресурса
            current_type = None
            for r in t.rows[2:]:
                vals = first_meaningful(r)
                if not vals:
                    continue
                joined = " ".join(vals)
                # Маркеры разделов
                resource_markers = ["ГВС","ХВС","Электроэнергия","Газ","Подтопления",
                                    "Экология","Несанкционированные свалки","Канализация","Другое"]
                if len(vals) == 1 and vals[0] in resource_markers:
                    current_type = vals[0]
                    continue
                if "Технологические нарушения" in joined or "Дата и время обращения" in joined:
                    continue
                if "Обращения на телефон" in joined or "Горячая линия» Главы" in joined:
                    break
                # строка данных: дата + описание (+ срок)
                if re.match(r'\d{2}\.\d{2}\.\d{4}', vals[0]) or (current_type and len(vals) >= 2):
                    date = vals[0] if re.match(r'\d{2}\.\d{2}\.\d{4}', vals[0]) else ""
                    # описание — самое длинное значение
                    desc = max(vals, key=len) if vals else ""
                    term = vals[-1] if len(vals) > 1 and vals[-1] != desc and len(vals[-1]) < 60 else ""
                    # Фильтр: справочная/консультативная информация — это НЕ технологическое
                    # нарушение (напр. "в адрес ЕДДС поступило N сообщений справочного и
                    # консультативного характера"). Складываем такие записи отдельно.
                    low = desc.lower()
                    is_info = (
                        ("справочн" in low and "консультатив" in low)
                        or ("справочного и консультативного характера" in low)
                        or ("поступило" in low and "сообщени" in low and "характера" in low)
                    )
                    if is_info:
                        data.setdefault("info_messages", []).append({
                            "datetime": date,
                            "desc": desc,
                        })
                        continue
                    data["tech_violations"].append({
                        "resource": current_type or "",
                        "datetime": date,
                        "desc": desc,
                        "term": term,
                    })

            # Категории горячей линии Главы (после маркера "Обращения на телефон")
            collecting = False
            for r in t.rows:
                vals = first_meaningful(r)
                joined = " ".join(vals)
                if "Обращения на телефон" in joined and "Главы" in joined:
                    collecting = True
                    continue
                if collecting:
                    if "поступило" in joined and "обращений" in joined:
                        data["appeals"]["hotline_total_text"] = joined
                        continue
                    # category | number
                    cv = cell_texts_dedup(r)
                    if len(cv) >= 2 and cv[-1] not in ('',):
                        cat = cv[0].rstrip(':')
                        num = cv[-1]
                        if cat and cat != num:
                            data["hotline_categories"].append({"category": cat, "count": num})
            break

    # --- Статистика вызовов ЕДДС ---
    for t in tables:
        if t.rows and "статистика поступивших вызовов" in clean(t.rows[0].cells[0].text).lower():
            labels = cell_texts_dedup(t.rows[1]) if len(t.rows) > 1 else []
            values = cell_texts_dedup(t.rows[2]) if len(t.rows) > 2 else []
            data["edds_stats"] = dict(zip(labels, values))
            break

    # --- Статистика выездов СОЛН СПАС ---
    for t in tables:
        if t.rows and "всего выездов" in clean(t.rows[0].cells[0].text).lower():
            labels = cell_texts_dedup(t.rows[0])
            values = cell_texts_dedup(t.rows[1]) if len(t.rows) > 1 else []
            data["spas_stats"] = dict(zip(labels, values))
            break

    # --- Уровни воды ---
    for t in tables:
        rows = t.rows
        if rows and "0" in clean(rows[0].cells[0].text) and "поста" in "".join(clean(c.text) for c in rows[0].cells):
            pass
    for t in tables:
        txt0 = clean(t.rows[0].cells[0].text) if t.rows else ""
        if ("река" in txt0.lower() or "поста" in "".join(clean(c.text) for r in t.rows for c in r.cells).lower()) and len(t.rows) >= 2:
            # парами: имя поста / значения
            i = 0
            rws = t.rows
            while i < len(rws):
                name_vals = first_meaningful(rws[i])
                name = name_vals[0] if name_vals else ""
                if "река" in name.lower() and i + 1 < len(rws):
                    val_vals = first_meaningful(rws[i+1])
                    zero = next((v for v in val_vals if "0" in v and "поста" in v), "")
                    actual = next((v for v in val_vals if "ктуальн" in v), "")
                    data["water_levels"].append({
                        "post": name,
                        "zero": clean(zero),
                        "actual": clean(actual),
                    })
                    i += 2
                else:
                    i += 1
            if data["water_levels"]:
                break

    # --- Погода (последняя таблица) ---
    for t in tables:
        head = [clean(c.text) for c in t.rows[0].cells] if t.rows else []
        if any("Температура" in h for h in head):
            labels = cell_texts_dedup(t.rows[0])
            values = cell_texts_dedup(t.rows[1]) if len(t.rows) > 1 else []
            data["weather"]["forecast"] = dict(zip(labels, values))
            break

    # --- Обращения по КПО "Нева" (неприятный/свалочный запах) ---
    # Считаем ТОЛЬКО детальные обращения раздела ЕДДС (технологические нарушения,
    # подраздел "Экология"), отфильтрованные по упоминанию "Нева".
    # Формулировки плавают: "Неприятный запах с КПО Нева", "Запах с КПО Нева",
    # "Свалочный запах с КПО Нева" — ловим просто по слову "нева" + "запах".
    neva_items = []
    for v in data.get("tech_violations", []):
        desc = (v.get("desc") or "")
        low = desc.lower()
        if "нева" in low and "запах" in low:
            neva_items.append({
                "datetime": v.get("datetime", ""),
                "desc": desc,
                "term": v.get("term", ""),
            })
    # Справочно: категория "Экология" в блоке "Горячая линия Главы" (агрегат, без
    # детализации). НЕ приравниваем к "Нева" — храним отдельно как контекст.
    hotline_eco = None
    for c in data.get("hotline_categories", []):
        if c.get("category", "").strip().lower().startswith("эколог"):
            raw = str(c.get("count", "")).strip()
            try:
                hotline_eco = int(re.sub(r"[^\d]", "", raw)) if re.search(r"\d", raw) else 0
            except Exception:
                hotline_eco = 0
            break
    data["neva"] = {
        "edds_count": len(neva_items),       # точное число поимённых обращений на ЕДДС
        "edds_items": neva_items,            # сами обращения (с маскированием на этапе рендера)
        "hotline_eco_count": hotline_eco,    # справочно: "Экология" горячей линии Главы
    }

    # --- Детекция смертельных случаев (подсветка красным) ---
    # Ищем признаки гибели людей во всех текстах: происшествия (вкл. ДТП),
    # их события и технологические нарушения. Словоформы разные, поэтому
    # ловим по корням. Исключаем «без пострадавших / погибших нет».
    FATAL_RE = re.compile(
        r"(?:погиб|погибш|смертельн|со\s*смертельн|летальн|скончал|"
        r"труп|обнаружен[аоы]?\s+мёртв|обнаружен[аоы]?\s+мертв|мёртв|мертв|"
        r"самоубийств|убийств|утону|гибел[ьи])",
        re.IGNORECASE,
    )
    # фразы-опровержения, при которых совпадение НЕ считается летальным
    NEG_RE = re.compile(
        r"(?:погибш\w*\s+нет|без\s+погибш|жертв\s+нет|без\s+жертв|"
        r"пострадавш\w*\s+нет|без\s+пострадавш|никто\s+не\s+погиб)",
        re.IGNORECASE,
    )

    def _is_fatal(text):
        if not text:
            return False
        if NEG_RE.search(text):
            return False
        return bool(FATAL_RE.search(text))

    def _count_val(raw):
        """Счётчик категории в число. 'нет'/'-'/'' → 0."""
        s = str(raw or "").strip().lower()
        if s in ("", "-", "–", "нет", "отсутствует", "0"):
            return 0
        m = re.search(r"\d+", s)
        return int(m.group()) if m else 0

    # Категории-счётчики, означающие гибель людей (по точному имени типа)
    FATAL_COUNTERS = ("погибло", "утонуло", "убийств", "самоубийств")

    fatal_items = []
    seen_texts = set()

    def _add(section, itype, text, **extra):
        key = (section, (text or "").strip()[:80])
        if key in seen_texts:
            return
        seen_texts.add(key)
        rec = {"section": section, "type": itype, "text": text}
        rec.update(extra)
        fatal_items.append(rec)

    # 1) происшествия: (а) счётчики «Погибло/Утонуло/Убийство/Самоубийства» > 0;
    #    (б) текстовые упоминания гибели в описаниях событий.
    for inc in data.get("incidents", []):
        itype = (inc.get("type", "") or "").strip()
        low_t = itype.lower()
        evs = inc.get("events") or []
        # (а) счётчики гибели
        if any(low_t.startswith(c) for c in FATAL_COUNTERS):
            n = _count_val(inc.get("count"))
            if n > 0:
                detail = "; ".join(e for e in evs if e) if evs else ""
                _add("incidents", itype, f"{itype}: {n}" + (f" — {detail}" if detail else ""),
                     count=n, counter=True)
            continue
        # (б) текстовые описания (пропускаем голые «См. раздел …»)
        for ev in evs:
            if "см. раздел" in ev.lower():
                continue
            if _is_fatal(ev):
                _add("incidents", itype, ev)
    # 2) технологические нарушения
    for v in data.get("tech_violations", []):
        desc = v.get("desc") or ""
        if _is_fatal(desc):
            _add("tech", v.get("resource", ""), desc, datetime=v.get("datetime", ""))

    data["fatal"] = {
        "has": len(fatal_items) > 0,
        "count": len(fatal_items),
        "items": fatal_items,
    }

    return data


if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument("path")
    ap.add_argument("-o", "--out", default=None)
    a = ap.parse_args()
    out = parse(a.path)
    js = json.dumps(out, ensure_ascii=False, indent=2)
    if a.out:
        with open(a.out, "w", encoding="utf-8") as f:
            f.write(js)
        print(f"OK -> {a.out} (дата {out.get('meta',{}).get('date')})")
    else:
        print(js)
